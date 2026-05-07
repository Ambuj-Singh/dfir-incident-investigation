#!/usr/bin/env python3
"""Alternate converter using python-docx for DOCX and fpdf for PDF generation."""
import sys
from pathlib import Path


def make_docx(md_path: Path, out_path: Path):
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    text = md_path.read_text(encoding='utf-8')
    for line in text.splitlines():
        line = line.rstrip()
        if line.startswith('#'):
            level = line.count('#', 0, line.find(' ')) if ' ' in line else 1
            para = doc.add_paragraph()
            run = para.add_run(line.lstrip('# ').strip())
            run.bold = True
            run.font.size = Pt(max(14 - (level - 1) * 2, 10))
        else:
            doc.add_paragraph(line)
    doc.save(str(out_path))


def make_pdf(md_path: Path, out_path: Path):
    from fpdf import FPDF

    text = md_path.read_text(encoding='utf-8')
    # normalize common unicode characters to ASCII for core fonts
    replacements = {
        '\u2014': '--',  # em dash
        '\u2013': '-',   # en dash
        '\u2018': "'",  # left single quote
        '\u2019': "'",  # right single quote
        '\u201c': '"',  # left double quote
        '\u201d': '"',  # right double quote
        '\u2026': '...', # ellipsis
    }
    for k, v in replacements.items():
        text = text.replace(k, v)
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font('Arial', size=10)
    for para in text.splitlines():
        if para.strip() == '':
            pdf.ln(4)
            continue
        # ensure long unbroken sequences are chunked to avoid FPDF errors
        words = para.split(' ')
        safe_words = []
        for w in words:
            if len(w) > 90:
                # break into 80-char chunks
                parts = [w[i:i+80] for i in range(0, len(w), 80)]
                safe_words.extend(parts)
            else:
                safe_words.append(w)
        safe_para = ' '.join(safe_words)
        try:
            pdf.multi_cell(0, 5, safe_para)
        except Exception:
            # final fallback: reduce to ASCII-only
            safe = safe_para.encode('ascii', errors='replace').decode('ascii')
            pdf.multi_cell(0, 5, safe)
    pdf.output(str(out_path))


def main():
    md = Path('docs/dfir_handbook_chapters.md')
    if not md.exists():
        print('Source markdown not found:', md)
        sys.exit(1)
    out_docx = Path('docs/dfir_handbook_chapters.docx')
    out_pdf = Path('docs/dfir_handbook_chapters.pdf')
    make_docx(md, out_docx)
    print('Wrote', out_docx)
    make_pdf(md, out_pdf)
    print('Wrote', out_pdf)


if __name__ == '__main__':
    main()
