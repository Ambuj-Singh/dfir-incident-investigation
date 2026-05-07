#!/usr/bin/env python3
"""Simple converter: Markdown -> DOCX (basic) and PDF (plain text)

This script performs a minimal conversion: it reads the Markdown file,
preserves headings as bold lines in the DOCX, and writes the raw text
into a straightforward PDF layout. It's intended for generating a
downloadable artifact for reviewers; for fully formatted conversions,
use `pandoc` with a LaTeX backend.
"""
import sys
from pathlib import Path


def make_docx(md_path: Path, out_path: Path):
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    text = md_path.read_text(encoding='utf-8')
    for line in text.splitlines():
        line = line.rstrip()
        if line.startswith('#'):
            # simple heading handling
            level = line.count('#', 0, line.find(' ')) if ' ' in line else 1
            para = doc.add_paragraph()
            run = para.add_run(line.lstrip('# ').strip())
            run.bold = True
            run.font.size = Pt(max(14 - (level - 1) * 2, 10))
        else:
            doc.add_paragraph(line)
    doc.save(str(out_path))


def make_pdf(md_path: Path, out_path: Path):
    from reportlab.lib.pagesizes import LETTER
    from reportlab.pdfgen import canvas
    from reportlab.lib.units import inch

    text = md_path.read_text(encoding='utf-8')
    c = canvas.Canvas(str(out_path), pagesize=LETTER)
    width, height = LETTER
    margin = 0.75 * inch
    y = height - margin
    max_width = width - 2 * margin
    lines = []
    # naive wrap at 95 characters
    for para in text.splitlines():
        if not para:
            lines.append('')
            continue
        while len(para) > 0:
            chunk = para[:95]
            lines.append(chunk)
            para = para[95:]

    c.setFont('Helvetica', 10)
    for line in lines:
        if y < margin:
            c.showPage()
            y = height - margin
            c.setFont('Helvetica', 10)
        c.drawString(margin, y, line)
        y -= 12
    c.save()


def main():
    md = Path('docs/dfir_handbook_chapters.md')
    if not md.exists():
        print('Source markdown not found:', md)
        sys.exit(1)
    out_docx = Path('docs/dfir_handbook_chapters.docx')
    out_pdf = Path('docs/dfir_handbook_chapters.pdf')
    make_docx(md, out_docx)
    print('Wrote', out_docx)
    make_pdf(md, out_pdf)
    print('Wrote', out_pdf)


if __name__ == '__main__':
    main()
